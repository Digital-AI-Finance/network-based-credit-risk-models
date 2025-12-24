#!/usr/bin/env python3
"""
Fill SNSF Word template with Final Scientific Report content.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

# Paths
TEMPLATE_PATH = Path(__file__).parent.parent / "final_scientific_report" / "ScientificReport_Osterrieder.docx"
OUTPUT_PATH = Path(__file__).parent.parent / "SNSF_Report_205487.docx"

# Content for the report
HEADER_INFO = {
    "Name": "Prof. Dr. Joerg Osterrieder",
    "Project number": "205487"
}

SECTION_1_1 = """Main Objective: To advance our understanding of credit risk modeling in P2P lending markets by designing and empirically verifying new network-based credit risk models. This project has attributes of a methodological and empirical project with practical impact, addressing information asymmetry inherent in P2P lending through network analysis. The main objective has been fully achieved through five distinct contributions that delivered novel methodology, rigorous empirical validation, and practical tools for interpretable credit risk assessment.

Contribution 1 (Methodological): Supervised Network-Based Credit Risk Models

Previous methods in the literature (Ahelegbey et al., 2019; Giudici et al., 2019, 2020) ignored loan status when constructing borrower networks, leading to unsupervised network-based learning approaches. Our approach fundamentally differs by utilizing class information (loan default status) to construct supervised networks, recognizing that supervised learning algorithms generally outperform unsupervised approaches for prediction tasks. We developed a two-step machine learning methodology that first constructs borrower similarity graphs using mixed-type distance measures (Gower distance for handling both continuous and categorical borrower attributes) and then extracts multiple network centrality measures as predictive features. The centrality measures include PageRank, betweenness centrality, closeness centrality, Katz centrality, and hub/authority scores, each capturing different aspects of a borrower's position within the similarity network. Graph construction employed k-nearest-neighbor connection rules combined with similarity threshold-based edge selection to ensure network connectivity while maintaining meaningful structure. This contribution has been fully achieved. Evidence: Liu et al. (2024), "Leveraging network topology for credit risk assessment in P2P lending: A comparative study under the lens of machine learning," Expert Systems with Applications, Vol. 252, Article 124100 (DOI: 10.1016/j.eswa.2024.124100), with 17 citations in Scopus demonstrating rapid adoption by the research community.

Contribution 2 (Methodological): Cross-Validation for Network Hyperparameters

In contrast to previous studies that used fixed, ad-hoc parameter choices, our work explicitly acknowledges that network creation and feature extraction depend on a set of hyperparameters whose optimal values vary depending on the specific dataset and network construction approach. We implemented systematic cross-validation to tune network hyperparameters and resulting features. The hyperparameter search space included similarity thresholds ranging from 0.1 to 0.9, k values for k-nearest-neighbor graphs ranging from 5 to 50, and different configurations of centrality measure computation. We employed 5-fold cross-validation throughout the model development process to ensure that hyperparameter selection was not overfit to a specific data split. This systematic approach to hyperparameter optimization represents a methodological advance over prior work and ensures reproducibility across different P2P lending contexts. This contribution has been fully achieved. Evidence: Methods sections in Liu et al. (2024) Expert Systems with Applications and Finance Research Letters publications document the complete hyperparameter tuning methodology; Zenodo deposits provide code for reproducible hyperparameter sensitivity analysis.

Contribution 3 (Methodological): Multiple Networks with Bootstrap Aggregation

Previous studies in the literature created only one network representation. We designed methods to create multiple networks that differ in two key aspects: (i) they utilize different random subsets of variables for similarity computation, and (ii) they rely on bootstrap aggregation (bagging) to combine predictions from models trained on different network representations. This approach directly addresses data noisiness, which is ignored in the existing literature on network-based credit scoring. The ensemble methodology was validated using four machine learning models: Elastic Net regression, Random Forest, Multi-Layer Perceptron (MLP) neural networks, and XGBoost gradient boosting. To verify that network features contain genuine predictive signal rather than spurious correlations, we conducted robustness checks using shuffled centrality features where network positions were randomly permuted across borrowers. When centrality features were shuffled, predictive power dropped substantially, confirming that the network topology captures genuine information about default risk rather than acting as a proxy for other features. This contribution has been fully achieved. Evidence: Ensemble methodology and robustness validation results are documented in the published papers and reproducible through Zenodo deposit 17991107.

Contribution 4 (Empirical): Validation Across Multiple P2P Datasets

We enriched the empirical literature, where most studies have used fewer than two P2P market datasets, by validating our models across different market platforms to observe credit drivers and method usefulness across diverse contexts. Primary empirical analysis used the Bondora dataset, a leading European P2P lending platform, containing 231,039 individual borrowers characterized through 112 categorical and continuous variables, with loan origination dates spanning from June 16, 2009 to April 21, 2022. The dataset includes detailed borrower demographics, financial attributes, and past credit market interactions. Comparative validation using LendingClub data from the United States market confirmed cross-market validity of our methodology. This geographic and regulatory diversity (European vs. US markets) strengthens the generalizability of our findings. A notable empirical finding was that degree centrality alone provides substantial predictive power for default risk, sometimes matching or exceeding more complex centrality measures like PageRank or betweenness centrality, suggesting that simpler network metrics may suffice for practical credit risk applications. This contribution has been fully achieved. Evidence: Liu et al. (2024), "Network centrality and credit risk: A comprehensive analysis of peer-to-peer lending dynamics," Finance Research Letters, Vol. 63, Article 105308 (DOI: 10.1016/j.frl.2024.105308), with 11 citations.

Contribution 5 (Practical): Explainable AI for Interpretable Credit Risk Models

We investigated the applicability of existing XAI methods to credit scoring models, enabling the development of interpretable credit risk models that address regulatory requirements for explainable automated lending decisions. SHAP (SHapley Additive exPlanations) values were computed for all features including network centrality measures, providing both global feature importance (expected effect of each variable on the outcome across all loans) and local explanations (expected effect of each variable for a specific individual loan application). LIME (Local Interpretable Model-agnostic Explanations) was implemented for complementary local interpretability. Beyond standard XAI techniques, we developed surrogate decision trees for regime detection model interpretability and a manual tree-based interpretation framework for understanding macro-regime classification in credit markets. These interpretability tools address the regulatory requirements emphasized by policymakers concerning automatized lending solutions, including GDPR requirements for explanation of automated decisions. This contribution has been fully achieved. Evidence: SHAP explainability notebooks in Zenodo deposit 17991107 (journal article reproducibility); tree-based interpretation framework in Zenodo deposit 17990398 (state-dependent pricing interpretability); all 12 Zenodo deposits ensure complete reproducibility of interpretability analyses.

Additional Project Outcomes

The project trained two PhD researchers, Lennart John Baals and Yiting Liu, who developed expertise in digital finance, network modeling, machine learning, and explainable AI. Both researchers are progressing toward completion of their doctoral theses at the University of Twente in collaboration with Bern University of Applied Sciences. We established substantive research collaborations with five international institutions across four continents: Masaryk University (Czech Republic) for joint publications and in-depth methodological exchanges; Columbia University (USA) for method exchanges and personnel exchange visits; American University of Sharjah (UAE) for joint publications and personnel exchange; Renmin University of China (China) for method exchanges and collaborative publications; and University of Manchester (UK) for joint publications. The PI serves as Action Chair of COST Action CA19130 (Fintech and AI in Finance), coordinating research collaboration across 40+ European institutions, and as Coordinator of the MSCA Industrial Doctoral Network on Digital Finance. Additional funding secured totaling CHF 90,000 includes two SNSF Mobility Grants (CHF 20,000 each) supporting researcher exchanges and a Leading House Asia grant (CHF 50,000) for related digital assets research."""

SECTION_1_2 = """Data Access Challenges: The original proposal planned validation across seven P2P platforms: Lending Club, Prosper, Zopa, Mintos, Bondora, Home Credit, and Kiva. Due to evolving data access restrictions, platform terms of service, and proprietary data policies, the scope of empirical validation was refined during project execution. Bondora's terms of service (Section 13.4) prevent redistribution of downloaded data, limiting our ability to share raw datasets while maintaining compliance. Similar restrictions applied to other planned platforms. We therefore focused primary empirical analysis on Bondora (European market) with comparative validation using LendingClub (US market), providing both geographic and regulatory diversity. This constraint led us to emphasize reproducible code and comprehensive documentation in our Zenodo deposits, ensuring that researchers with legitimate data access can replicate our methodology.

Platform Evolution: The P2P lending landscape evolved significantly during the project period (October 2022 - August 2025). Several platforms reduced operations, changed business models, or implemented more restrictive data access policies. The global rise in interest rates during 2022-2023 affected P2P lending volumes and platform viability across markets. This required flexibility in our empirical strategy, ultimately strengthening our focus on methodological contributions that remain applicable across different platforms regardless of specific data availability.

Computational Scaling: Network construction for large loan portfolios presented computational challenges. Building similarity graphs for 231,039 borrowers requires efficient algorithms to avoid quadratic complexity in pairwise distance computation. We addressed this through efficient similarity thresholding (computing distances only for potentially connected pairs) and k-nearest-neighbor connection rules that maintain network connectivity while managing computational load. These algorithmic choices are documented in our code repositories for researchers facing similar scaling challenges.

Unexpected Finding: Our empirical analysis revealed that network position measured by simple degree centrality alone provides substantial predictive power for default risk, sometimes matching or exceeding more computationally expensive centrality measures such as PageRank, betweenness, or Katz centrality. This finding has practical implications for platform operators seeking interpretable and computationally efficient risk indicators. The simplicity of degree centrality (counting direct connections) makes it easier to explain to stakeholders and regulators compared to more complex network metrics, supporting adoption of network-based credit assessment in practice."""

SECTION_1_3 = """Methodological Contribution: We established network topology as a viable and valuable feature source for credit risk modeling in P2P lending, demonstrating that borrower position within similarity-based networks contains predictive information about default risk beyond what is captured by traditional borrower-level attributes alone. Our two-step machine learning approach (network construction followed by centrality-based feature extraction) provides a replicable framework for combining structural and attribute-based predictors. The systematic literature review conducted as part of this project analyzed 78 articles selected from 1,066 initially retrieved records through structured screening in Scopus and Web of Science, employing double-blind coding to ensure consistency. This review provides a comprehensive synthesis of graph-based credit modeling approaches across P2P lending and banking network applications, identifying research gaps and methodological best practices.

Practical Impact: Our models and code are directly applicable to P2P lending platforms for credit risk assessment. The interpretability framework using SHAP and LIME enables platform operators to explain credit decisions to borrowers and regulators, addressing transparency requirements in automated lending decisions mandated by regulations including GDPR. The finding that simple degree centrality provides substantial predictive power offers practitioners an accessible entry point to network-based credit modeling without requiring expertise in complex network science methods.

Open Science Contribution: All 12 Zenodo deposits are 100% open access, ensuring complete reproducibility of our findings. The deposits include LaTeX manuscript sources, Jupyter notebooks implementing the complete analysis pipeline from data preprocessing through model training and evaluation, Python and R scripts for network construction and centrality computation, and comprehensive documentation enabling other researchers to validate, extend, and build upon our work. The curated Bondora dataset is archived at the Open Science Framework (OSF) at https://osf.io/jnpfs/.

Capacity Building: Two PhD researchers were trained in cutting-edge methods at the intersection of network science, machine learning, and finance, developing skills directly relevant to the growing fintech sector. Knowledge transfer activities included presentations at international research events in eight countries: Switzerland, Germany, Turkey, Netherlands, Hong Kong, China, Czech Republic, and UAE.

Research Network Leadership: The PI's roles as COST Action CA19130 Chair and MSCA Digital Finance Network Coordinator enabled broad dissemination of project findings across 40+ European institutions, creating lasting infrastructure for fintech research collaboration that extends well beyond this individual project. The COST Action facilitated policy discussions at EU level including events in Brussels addressing AI in finance implications for financial regulation."""

EXTENDED_INFO = """Extended Project Information

Project Statistics: Eight core publications with 28+ total citations received, 12 Zenodo deposits ensuring reproducibility, 12 conference presentations across 8 countries, 2 PhD researchers trained, and CHF 90,000 in additional funding secured.

1.4 Impact Statement

Scientific Impact: Our publications have received 28+ citations within 12 months of publication, indicating rapid adoption by the research community. The two-step machine learning methodology combining network centrality with traditional credit features has been referenced in subsequent studies on P2P lending risk assessment, and the systematic literature review provides a foundational reference for researchers entering the field of graph-based credit modeling.

Economic Impact: The developed models and open-source code are directly applicable to P2P lending platforms for improved credit risk assessment. By enabling more accurate default prediction, platforms can better price loans to reflect true risk, reduce losses from defaults, and offer more competitive rates to creditworthy borrowers. The interpretability framework addresses regulatory requirements, reducing compliance costs for platforms adopting automated credit decisions.

Social Impact: Improved credit risk models contribute to financial inclusion by enabling P2P platforms to serve borrowers who may be underserved by traditional banking. More accurate risk assessment reduces adverse selection problems, protecting retail investors who fund P2P loans from excessive default losses. The transparency framework enhances borrower trust in automated credit decisions by providing explanations for lending outcomes.

Policy Impact: The PI's leadership of COST Action CA19130 facilitated policy discussions at EU level, including events in Brussels addressing AI in finance policy implications and contributing to the broader discourse on responsible AI adoption in financial services.

Educational Impact: The project trained two PhD researchers in cutting-edge methods at the intersection of network science, machine learning, and finance. The publication "Towards a new PhD Curriculum for Digital Finance" (Open Research Europe, 2024, DOI: 10.12688/openreseurope.16513.1) disseminates best practices for doctoral training in this emerging field, contributing to curriculum development beyond this specific project.

1.5 Sustainability Plan

Data Preservation: Twelve Zenodo deposits are archived with persistent DOIs ensuring permanent accessibility and citability through CERN's infrastructure. The curated Bondora P2P lending dataset is archived at the Open Science Framework (OSF). Code repositories are maintained under the Digital-AI-Finance organization on GitHub. All outputs are released under Creative Commons Attribution 4.0 (CC-BY 4.0) licensing, enabling unrestricted reuse with attribution.

Code Maintainability: All Jupyter notebooks and Python/R scripts include dependency specifications (requirements.txt, environment files) enabling reproduction with specified package versions. Reproducibility has been verified through independent testing. Documentation is embedded in code through comments and supplementary README files.

Knowledge Transfer Continuation: COST Action CA19130 (Fintech and AI in Finance) continues beyond project end with the PI serving as Action Chair. The MSCA Industrial Doctoral Network on Digital Finance continues training next-generation researchers with the PI as Coordinator. Digital finance research continues at Bern University of Applied Sciences building on this project's foundations.

Appendix A: Peer-Reviewed Publications

Liu, Y., Baals, L.J., Osterrieder, J., Hadji-Misheva, B. (2024). Leveraging network topology for credit risk assessment in P2P lending: A comparative study under the lens of machine learning. Expert Systems with Applications, 252(B), 124100. DOI: 10.1016/j.eswa.2024.124100. 17 citations.

Liu, Y., Baals, L.J., Osterrieder, J., Hadji-Misheva, B. (2024). Network centrality and credit risk: A comprehensive analysis of peer-to-peer lending dynamics. Finance Research Letters, 63, 105308. DOI: 10.1016/j.frl.2024.105308. 11 citations.

Baumohl, E., Lyocsa, S., Vasanicova, P. (2024). Macroeconomic environment and the future performance of loans: Evidence from three peer-to-peer platforms. International Review of Financial Analysis, 95, 103416. DOI: 10.1016/j.irfa.2024.103416.

Baals, L.J., Osterrieder, J., Hadji-Misheva, B., Liu, Y. (2024). Towards a new PhD Curriculum for Digital Finance. Open Research Europe, 4, 16513. DOI: 10.12688/openreseurope.16513.1.

Lyocsa, S., Todorovic, N., Baals, L.J., Gao, H. (2025). Alpha-threshold networks in credit risk models. Quantitative Finance. DOI: 10.1080/14697688.2025.2465697.

Lyocsa, S., Plat, V., Gao, H. (2025). A fuzzy framework for realized volatility prediction. Journal of Forecasting. DOI: 10.1002/for.70082.

Submitted: Baals, L.J., et al. (2025). Network Evidence on Credit-Risk Pricing in P2P Lending. SSRN 5276337. Baals, L.J., et al. (2025). State-Dependent Pricing in FinTech Credit: Evidence from P2P Lending. SSRN 5421207.

Main Publication Output - PhD Researchers (ORCID)

The following publications are registered in ORCID for the project's PhD researchers:

Lennart John Baals (ORCID: 0000-0002-7737-9675):
- Baals, L.J., Liu, Y. et al. (2024). "Leveraging network topology for credit risk assessment in P2P lending: A comparative study under the lens of machine learning". Expert Systems with Applications, 252, 124100. DOI: https://doi.org/10.1016/j.eswa.2024.124100
- Baals, L.J., Liu, Y. et al. (2024). "Network centrality and credit risk: A comprehensive analysis of peer-to-peer lending dynamics". Finance Research Letters, 63, 105308. DOI: https://doi.org/10.1016/j.frl.2024.105308

Yiting Liu (ORCID: 0009-0006-9554-8205):
- Liu, Y., Baals, L.J. et al. (2024). "Leveraging network topology for credit risk assessment in P2P lending: A comparative study under the lens of machine learning". Expert Systems with Applications, 252, 124100. DOI: https://doi.org/10.1016/j.eswa.2024.124100
- Liu, Y., Baals, L.J. et al. (2024). "Network centrality and credit risk: A comprehensive analysis of peer-to-peer lending dynamics". Finance Research Letters, 63, 105308. DOI: https://doi.org/10.1016/j.frl.2024.105308
- Liu, Y. et al. (2024). "Towards a new PhD Curriculum for Digital Finance". Open Research Europe, 4, 16513. DOI: https://doi.org/10.12688/openreseurope.16513.1
- Liu, Y. et al. (2023). "Navigating the Environmental, Social, and Governance (ESG) landscape: constructing a robust and reliable scoring engine". Open Research Europe. DOI: https://doi.org/10.12688/openreseurope.16278.1

Appendix B: Open Science Deposits (Zenodo)

1. COST FinAI Meets Istanbul Conference Event May 20-21, 2024. Baals, Lennart John (2024). Presentation. https://zenodo.org/records/17964900

2. State-Dependent Pricing in FinTech Credit: Evidence from P2P Lending. Baals, Lennart John (2025). Working paper. https://zenodo.org/records/17990398

3. Network Evidence on Credit-Risk Pricing in P2P Lending. Baals, Lennart John (2025). Working paper. https://zenodo.org/records/17990873

4. Leveraging Network Topology for Credit Risk Assessment in P2P Lending: A Comparative Study under the Lens of Machine Learning. Baals, Lennart John (2025). Journal article. https://zenodo.org/records/17991107

5. PhD Qualifier Report and Presentation delivered by Lennart John Baals at the University of Twente. Baals, Lennart John (2025). Proposal. https://zenodo.org/records/17992215

6. A Systematic Literature Review on Graph-Based Models in Credit Risk Assessment. Baals, Lennart John (2025). Presentation. https://zenodo.org/records/17992322

7. Leveraging Network Topology for Credit Risk Assessment in P2P Lending (Bern Conference 2023). Baals, Lennart John (2025). Presentation. https://zenodo.org/records/17992484

8. Identifying Mispriced Loans through Interest Rate-Based Network Analysis and Clustering in P2P Lending Markets. Baals, Lennart John (2025). Presentation. https://zenodo.org/records/17992591

9. Data and Code to reproduce results in paper "Network centrality and credit risk". Liu, Yiting (2024). Computational notebook. https://zenodo.org/records/17989119

10. Data and Code to reproduce results in paper "Credit Risk Prediction via Graph Neural Networks with Homophily-Guided Graph Construction". Liu, Yiting (2026). Computational notebook. https://zenodo.org/records/17990002

11. Data and Code to reproduce results in paper "Explaining Regime Dynamics: A Tree-based Interpretation Framework for R2-RD Models". Liu, Yiting (2026). Computational notebook. https://zenodo.org/records/17990140

12. Data and Code to reproduce results in paper "Leveraging network topology for credit risk assessment in P2P lending". Liu, Yiting (2024). Computational notebook. https://zenodo.org/records/17990581

Appendix C: Academic Events

December 17, 2024: 4th International Symposium on Big Data and AI, Hong Kong - Systematic Literature Review on Graph-Based Credit Models

September 29, 2024: 8th Bern Conference on Fintech and AI in Finance, Bern, Switzerland

September 2024: AI Finance Insights: Pioneering the Future of Fintech, Istanbul, Turkey

May 20-21, 2024: COST FinAI Meets Istanbul Conference, Turkey

May 2024: COST FinAI PhD School, Treviso, Italy

May 2024: COST FinAI Brussels Conference, Belgium

December 2023: 16th ERCIM WG / 17th CFE Conference, Berlin, Germany

November 29, 2023: BFH Doctoral Seminar, Bern, Switzerland

September 27-29, 2023: 8th European COST Conference on AI in Finance, Bern, Switzerland

September 2023: European Summer School in Financial Mathematics, TU Delft, Netherlands

September 2023: Shenzhen Technology University International Week, China

June 2023: COST Action Training School, Enschede, Netherlands

Appendix D: Dataset

Bondora P2P Lending Dataset. Coverage: June 2009 - April 2022. Sample: 231,039 borrowers, 112 variables. DOI: 10.21227/33kz-0s65. License: CC-BY 4.0.

Appendix E: International Collaborations

Masaryk University (Czech Republic), Columbia University (USA), American University of Sharjah (UAE), Renmin University of China (China), University of Manchester (UK).

Appendix F: PhD Researchers

Lennart John Baals: PhD In Progress, BFH/University of Twente, Graph-based credit models and network analysis for credit risk assessment.

Yiting Liu: PhD In Progress, BFH/University of Twente, P2P lending risk modeling and network topology for credit risk.

Report submitted to: Swiss National Science Foundation (SNSF). Report date: December 2025. Data source: https://data.snf.ch/grants/grant/205487"""


def find_table_cell_after(doc, search_text):
    """Find the table cell that comes after the cell containing search_text."""
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if search_text in cell.text:
                    if i + 1 < len(row.cells):
                        return row.cells[i + 1]
    return None


def find_paragraph_containing(doc, search_text):
    """Find paragraph containing specific text."""
    for para in doc.paragraphs:
        if search_text in para.text:
            return para
    return None


def clear_and_add_content(doc, start_text, end_text, content):
    """Clear content between start and end markers, add new content."""
    start_idx = None
    end_idx = None

    for i, para in enumerate(doc.paragraphs):
        if start_text in para.text:
            start_idx = i
        if end_text and end_text in para.text and start_idx is not None:
            end_idx = i
            break

    if start_idx is None:
        print(f"Could not find: {start_text}")
        return

    # Find the next heading or end of document
    if end_idx is None:
        for i in range(start_idx + 1, len(doc.paragraphs)):
            style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
            if "Heading" in style or "1" in style:
                end_idx = i
                break
        if end_idx is None:
            end_idx = len(doc.paragraphs)

    # Clear existing content paragraphs (keep heading)
    paras_to_clear = []
    for i in range(start_idx + 1, end_idx):
        para = doc.paragraphs[i]
        # Skip if it's a heading
        style = para.style.name if para.style else ""
        if "Heading" not in style:
            paras_to_clear.append(para)

    for para in paras_to_clear:
        p = para._element
        p.getparent().remove(p)

    # Add new content after the heading
    heading_para = doc.paragraphs[start_idx]

    # Split content into paragraphs
    paragraphs = content.strip().split('\n\n')

    # Insert paragraphs after heading
    for para_text in paragraphs:
        if para_text.strip():
            new_para = doc.add_paragraph()
            # Move to correct position (after heading)
            heading_para._element.addnext(new_para._element)
            heading_para = new_para  # Update reference for next insertion

            # Add text with bold handling
            parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = new_para.add_run(part[2:-2])
                    run.bold = True
                else:
                    new_para.add_run(part)


def main():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    # Fill in Name
    name_cell = find_table_cell_after(doc, "Name")
    if name_cell:
        name_cell.text = HEADER_INFO["Name"]
        print("Filled: Name")

    # Fill in Project number
    project_cell = find_table_cell_after(doc, "Project")
    if project_cell:
        project_cell.text = HEADER_INFO["Project number"]
        print("Filled: Project number")

    # For sections, we need to find the italic instruction text and replace it
    # The template has italic instructions that need to be replaced with actual content

    # Find and replace section content
    in_section_1_1 = False
    in_section_1_2 = False
    in_section_1_3 = False

    paras_to_remove = []
    section_1_1_heading = None
    section_1_2_heading = None
    section_1_3_heading = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Detect section headings
        if "Achievement of research objectives" in text:
            in_section_1_1 = True
            in_section_1_2 = False
            in_section_1_3 = False
            section_1_1_heading = para
            continue
        elif "Challenges, negative results" in text or "Challenges," in text:
            in_section_1_1 = False
            in_section_1_2 = True
            in_section_1_3 = False
            section_1_2_heading = para
            continue
        elif "Contribution to knowledge" in text:
            in_section_1_1 = False
            in_section_1_2 = False
            in_section_1_3 = True
            section_1_3_heading = para
            continue

        # Mark instruction paragraphs for removal (italic text in sections)
        if in_section_1_1 or in_section_1_2 or in_section_1_3:
            if text and len(text) > 5:  # Non-empty paragraphs
                paras_to_remove.append(para)

    # Remove instruction paragraphs
    for para in paras_to_remove:
        p = para._element
        p.getparent().remove(p)
    print(f"Removed {len(paras_to_remove)} instruction paragraphs")

    # Add Section 1.1 content
    if section_1_1_heading:
        prev = section_1_1_heading
        for para_text in SECTION_1_1.strip().split('\n\n'):
            if para_text.strip():
                new_para = doc.add_paragraph()
                prev._element.addnext(new_para._element)
                prev = new_para

                # Check if it's a contribution heading
                if para_text.startswith("Contribution") or para_text.startswith("Main Objective") or para_text.startswith("Additional"):
                    run = new_para.add_run(para_text)
                    run.bold = True
                else:
                    new_para.add_run(para_text)
        print("Added Section 1.1 content")

    # Add Section 1.2 content
    if section_1_2_heading:
        prev = section_1_2_heading
        for para_text in SECTION_1_2.strip().split('\n\n'):
            if para_text.strip():
                new_para = doc.add_paragraph()
                prev._element.addnext(new_para._element)
                prev = new_para

                # Check if starts with bold heading pattern
                if para_text.split(':')[0] in ["Data Access Challenges", "Platform Evolution", "Computational Scaling", "Unexpected Finding"]:
                    parts = para_text.split(':', 1)
                    run = new_para.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        new_para.add_run(parts[1])
                else:
                    new_para.add_run(para_text)
        print("Added Section 1.2 content")

    # Add Section 1.3 content
    if section_1_3_heading:
        prev = section_1_3_heading
        for para_text in SECTION_1_3.strip().split('\n\n'):
            if para_text.strip():
                new_para = doc.add_paragraph()
                prev._element.addnext(new_para._element)
                prev = new_para

                # Check if starts with bold heading pattern
                headings = ["Methodological Contribution", "Practical Impact", "Open Science Contribution", "Capacity Building", "Research Network Leadership"]
                first_part = para_text.split(':')[0]
                if first_part in headings:
                    parts = para_text.split(':', 1)
                    run = new_para.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        new_para.add_run(parts[1])
                else:
                    new_para.add_run(para_text)
        print("Added Section 1.3 content")

    # Add Extended Project Information after Section 1.3
    if section_1_3_heading:
        # Find the last paragraph we added (prev still points to it)
        extended_headings = [
            "Extended Project Information", "Project Statistics", "1.4 Impact Statement",
            "Scientific Impact", "Economic Impact", "Social Impact", "Policy Impact", "Educational Impact",
            "1.5 Sustainability Plan", "Data Preservation", "Code Maintainability", "Knowledge Transfer Continuation",
            "Appendix A", "Main Publication Output", "Lennart John Baals", "Yiting Liu",
            "Appendix B", "Appendix C", "Appendix D", "Appendix E", "Appendix F",
            "Report submitted to", "The following publications"
        ]
        for para_text in EXTENDED_INFO.strip().split('\n\n'):
            if para_text.strip():
                new_para = doc.add_paragraph()
                prev._element.addnext(new_para._element)
                prev = new_para

                # Check if it's a heading or section title
                first_part = para_text.split(':')[0].strip()
                is_heading = any(h in first_part for h in extended_headings)

                if is_heading and ':' in para_text:
                    parts = para_text.split(':', 1)
                    run = new_para.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        new_para.add_run(parts[1])
                elif is_heading:
                    run = new_para.add_run(para_text)
                    run.bold = True
                else:
                    new_para.add_run(para_text)
        print("Added Extended Project Information")

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nSaved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
