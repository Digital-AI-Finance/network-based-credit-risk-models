#!/usr/bin/env python3
"""
Add PhD ORCID publications section to the existing ScientificReport_Osterrieder.docx
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path

INPUT_PATH = Path(__file__).parent.parent / "final_scientific_report" / "ScientificReport_Osterrieder.docx"
OUTPUT_PATH = Path(__file__).parent.parent / "SNSF_Report_205487.docx"

ORCID_SECTION = """Main Publication Output (ORCID 2024-2025)

The following publications are registered in ORCID for the project researchers (2024-2025):

Joerg Osterrieder (ORCID: 0000-0003-0189-8636):

1. "How can artificial intelligence help customer intelligence for credit portfolio management? A systematic literature review". International Journal of Information Management Data Insights. DOI: 10.1016/j.jjimei.2024.100234

2. "Stylized facts of metaverse non-fungible tokens". Physica A: Statistical Mechanics and its Applications. DOI: 10.1016/j.physa.2024.130103

3. "Leveraging network topology for credit risk assessment in P2P lending". Expert Systems with Applications. DOI: 10.1016/j.eswa.2024.124100 (with Baals, Liu)

4. "Network centrality and credit risk: A comprehensive analysis of peer-to-peer lending dynamics". Finance Research Letters. DOI: 10.1016/j.frl.2024.105308 (with Baals, Liu)

5. "Towards a new PhD Curriculum for Digital Finance". Open Research Europe. DOI: 10.12688/openreseurope.16513.1 (with Liu)

6. "Visual XAI tool". Zenodo. DOI: 10.5281/zenodo.10934115

7. "A discussion paper for possible approaches to building a statistically valid backtesting framework". SSRN. DOI: 10.2139/ssrn.4893677

8. "Enhancing Security in Blockchain Networks: Anomalies, Frauds, and Advanced Detection Techniques". arXiv. DOI: 10.48550/arxiv.2402.11231

9. "Ethical Artificial Intelligence, Fintech and Data Protection: A Path Forward for Training in Europe". SSRN. DOI: 10.2139/ssrn.4885037

10. "Forecasting Commercial Customers Credit Risk Through Early Warning Signals Data". SSRN. DOI: 10.2139/ssrn.4754568

11. "How can Consumers Without Credit History Benefit from Information Processing and Machine Learning Tools by Financial Institutions?". SSRN. DOI: 10.2139/ssrn.4730445

12. "Hypothesizing Multimodal Influence: Assessing the Impact of Textual and Non-Textual Data on Financial Instrument Pricing Using NLP and Generative AI". SSRN. DOI: 10.2139/ssrn.4698153

13. "Integrating Early Warning Systems with Customer Segmentation". SSRN. DOI: 10.2139/ssrn.4779632

14. "Integration of Early Warning Systems and Customer Segmentation Methods in the Financial Industry - A Systematic Literature Review". SSRN. DOI: 10.2139/ssrn.4730479

15. "Metaverse Non Fungible Tokens". SSRN. DOI: 10.2139/ssrn.4733153

16. "Modeling Commodity Price Co-Movement: Building on Traditional Methods & Exploring Applications of Machine Learning Models". SSRN. DOI: 10.2139/ssrn.4730474

17. "Predicting Retail Customers' Distress: Early Warning Systems and Machine Learning Applications". SSRN. DOI: 10.2139/ssrn.4730470

Lennart John Baals (ORCID: 0000-0002-7737-9675): Co-author on publications 3, 4 above.

Yiting Liu (ORCID: 0009-0006-9554-8205): Co-author on publications 3, 4, 5 above."""

def main():
    print(f"Loading: {INPUT_PATH}")
    doc = Document(INPUT_PATH)

    # Find "Appendix B" and insert ORCID section before it
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Appendix B" in para.text:
            insert_idx = i
            break

    if insert_idx is None:
        print("Could not find 'Appendix B' - appending to end")
        insert_idx = len(doc.paragraphs)

    # Add paragraphs in reverse order (since we're inserting before Appendix B)
    paragraphs = ORCID_SECTION.strip().split('\n\n')

    # Find the element before Appendix B
    appendix_b_para = doc.paragraphs[insert_idx]

    # Insert new paragraphs before Appendix B
    prev = None
    for para_text in paragraphs:
        if para_text.strip():
            new_para = doc.add_paragraph()

            # Move to correct position
            if prev is None:
                appendix_b_para._element.addprevious(new_para._element)
            else:
                prev._element.addnext(new_para._element)
            prev = new_para

            # Check if it's a heading
            if para_text.startswith("Main Publication Output") or \
               para_text.startswith("Lennart John Baals") or \
               para_text.startswith("Yiting Liu") or \
               para_text.startswith("The following"):
                run = new_para.add_run(para_text)
                if para_text.startswith("Main Publication Output"):
                    run.bold = True
                elif para_text.startswith("Lennart") or para_text.startswith("Yiting"):
                    run.bold = True
            else:
                new_para.add_run(para_text)

    print(f"Added ORCID publications section")
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
